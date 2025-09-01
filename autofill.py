#!/usr/bin/env python3
from PyQt6 import QtCore, QtWidgets, uic
import sys
import os
import docx
from datetime import datetime, timedelta
from threading import Thread
from time import sleep
from urllib.request import urlopen
import re

import sys
if sys.platform.startswith("win"):
    WIN = True
    WIN_css = 'font-size: 11pt;'
else:
    WIN = False
    WIN_css = ''

APP_PATH = os.path.abspath(os.path.dirname(sys.argv[0]))

# диапазон дат четвертей
list_date = ['01.09.2022', '29.10.2022', '07.11.2022', '30.12.2022', '09.01.2023', '25.03.2023', '03.04.2023', '25.05.2023', '06.02.2023', '12.02.2023']
list_date_holidays = ['23.02.2023', '24.02.2023', '08.03.2023', '01.05.2023', '08.05.2023', '09.05.2023']
dict_date_replaced = {'25.05.2023': 0}
range_date_holidays = []
list_days = []
all_days = 0
dict_days = {}
week_days = []
double_days = []
week_days2 = []
double_days2 = []
file_name = ''
count_hours = 0
table_number = -1
one = 2
column_with_days = 3
list_one = []
rewrite = False
diff_hours = False
klass1 = False
list_klass1 = []

pushButtonFill_text_dict = {
    True: 'Заполнить и перезаписать файл',
    False: 'Заполнить и создать копию'
}

dict_replaced_weekdays = {
    0: 'по понедельнику',
    1: 'по вторнику',
    2: 'по среде',
    3: 'по четвергу',
    4: 'по пятнице',
    5: 'по субботе'
}

list_weekdays = ['Понедельник', 'Вторник', 'Среда', 'Четверг', 'Пятница', 'Суббота']

version = 'Версия 1.8 от 29.08.2025'

def dateFmt(date):
    return(QtCore.QDate.fromString(date, 'dd.MM.yyyy'))

# загрузка окна
app = QtWidgets.QApplication(sys.argv)
win = uic.loadUi('main.ui')
win.groupBoxNewVer.hide()
win.tabWidget.tabBar().hide()
win.labelVersion.setText(version)
win_about = uic.loadUi('about.ui')
win_about.labelVersion.setText(version)
win_help = uic.loadUi('help.ui')
win_help.groupBoxLink.hide()
win_help.resize(500, 300)
win_dialog = uic.loadUi('dialog.ui')
win_dialog.comboBoxReplace.addItems(list_weekdays)
win_dialog.dateEditReplace.setDate(dateFmt(list_date[0]))

# получаем список всех виджетов и создаем переменные
for i in dir(win):
    try:
        exec('win.'+i+'.objectName()')
        exec('globals()[win.'+i+'.objectName()] = win.'+i)
    except:
        pass

def fixFontIfWIN():
    if WIN:
        font_size = 'QWidget {font-size: 11pt;}'
        win.setStyleSheet(font_size)
        win_dialog.setStyleSheet(font_size)
        win.pushButtonFill.setStyleSheet('QWidget {font-size: 16pt;}')
        for i in range(1, 5):
            globals()['label0'+str(i)].setStyleSheet('QWidget {font-size: 12pt;}')
        win_about.setStyleSheet(font_size)
        win_about.label.setStyleSheet('QWidget {font-size: 17pt;}')
        win_help.setStyleSheet(font_size)

def loadDateFromFile():
    global list_date
    if os.path.exists(os.path.join(APP_PATH, 'сохранённый диапазон дат.txt')):
        with open(os.path.join(APP_PATH, 'сохранённый диапазон дат.txt'), 'r') as f:
            list_date = []
            for line in f:
                list_date.append(line.strip())
        print('Dates Loaded from file')

def loadHolidaysFromFile():
    global list_date_holidays
    if os.path.exists(os.path.join(APP_PATH, 'сохранённые праздничные дни.txt')):
        with open(os.path.join(APP_PATH, 'сохранённые праздничные дни.txt'), 'r') as f:
            list_date_holidays = []
            for line in f:
                list_date_holidays.append(line.strip())
        print('Holidays loaded from file')

def loadHolidays2FromFile():
    global range_date_holidays
    if os.path.exists(os.path.join(APP_PATH, 'сохранённые дополнительные каникулы.txt')):
        with open(os.path.join(APP_PATH, 'сохранённые дополнительные каникулы.txt'), 'r') as f:
            range_date_holidays = []
            for line in f:
                range_date_holidays.append(line.split())
        print('Holidays2 loaded from file')

def loadReplacedFromFile():
    global dict_date_replaced
    if os.path.exists(os.path.join(APP_PATH, 'сохранённые замены дней недели.txt')):
        with open(os.path.join(APP_PATH, 'сохранённые замены дней недели.txt'), 'r') as f:
            dict_date_replaced = {}
            for line in f:
                line = line.strip()
                dict_date_replaced[line.split(';')[0]] = int(line.split(';')[1])
        print('Replaced loaded from file')

def loadDateToApp():
    for i in range(16):
        globals()['dateEdit'+str(i)].setDate(dateFmt(list_date[i]))

def loadHolidaysToApp():
    dateEditHoliday.setDate(dateFmt(list_date[0]))
    listWidgetHolidays.clear()
    listWidgetHolidays.addItems(list_date_holidays)

def loadHolidays2ToApp():
    if len(range_date_holidays):
        dateEditAdd1.setDate(dateFmt(range_date_holidays[-1][0]))
        dateEditAdd2.setDate(dateFmt(range_date_holidays[-1][1]))
    else:
        dateEditAdd1.setDate(dateFmt('01.09.2024'))
        dateEditAdd2.setDate(dateFmt('01.09.2024'))
    listWidgetHolidays2.clear()
    for el in range_date_holidays:
        item = QtWidgets.QListWidgetItem(el[0] + ' - ' + el[1])
        if el[2] == '1':
            item.setCheckState(QtCore.Qt.CheckState.Checked)
        else:
            item.setCheckState(QtCore.Qt.CheckState.Unchecked)
        listWidgetHolidays2.addItem(item)

def loadReplaceToApp():
    s = []
    for key, value in dict_date_replaced.items():
        s.append(key+' '+dict_replaced_weekdays[value])
    labelReplace.setText('Замена дней недели:\n'+'\n'.join(s))
    win_dialog.listWidgetReplace.clear()
    win_dialog.listWidgetReplace.addItems(s)

def connectSignals():
    for i in range(16):
        globals()['dateEdit'+str(i)].dateChanged.connect(lambda: changedDateEdit())
    for i in range(6):
        globals()['checkBoxWeek'+str(i)].clicked.connect(lambda: readWeekDays())
        globals()['spinBoxWeek'+str(i)].valueChanged.connect(lambda: readWeekDays())
        globals()['checkBoxWeek'+str(i)+'_2'].clicked.connect(lambda: readWeekDays())
        globals()['spinBoxWeek'+str(i)+'_2'].valueChanged.connect(lambda: readWeekDays())
    checkBoxYear.clicked.connect(lambda: readWeekDays())
    pushButton.clicked.connect(lambda: openFiles())
    pushButtonFill.clicked.connect(lambda: fill())
    frame_2.setStyleSheet('QWidget {}')
    frame_3.setStyleSheet('QWidget {}')
    comboBoxColumns.currentIndexChanged.connect(setColumnWithDay)
    comboBoxTables.textActivated.connect(setTable)
    action_1.triggered.connect(lambda: win_help.show())
    action_2.triggered.connect(lambda: win_about.show())
    win_help.pushButtonChekNewVersion.clicked.connect(lambda: buttonCheckNewVersionClick())
    checkBoxSetCol.clicked.connect(lambda: groupBoxSetCol.setStyleSheet('QWidget {}'))
    checkBoxRewrite.clicked.connect(checkBoxRewriteClick)
    checkBoxDiffHours.clicked.connect(checkBoxDiffHoursClick)
    pushButtonAddHoliday.clicked.connect(lambda: pushButtonAddHolidayClick())
    pushButtonDeleteHoliday.clicked.connect(lambda: pushButtonDeleteHolidayClick())
    pushButtonDeleteHoliday2.clicked.connect(lambda: pushButtonDeleteHoliday2Click())
    pushButtonReplace.clicked.connect(lambda: win_dialog.show())
    win_dialog.pushButtonAddReplace.clicked.connect(lambda: pushButtonAddReplaceClick())
    pushButtonAddHoliday2.clicked.connect(lambda: pushButtonAddHolidays2Click())
    win_dialog.pushButtonDeleteReplace.clicked.connect(lambda: pushButtonDeleteReplaceClick())
    checkBox1klass.clicked.connect(checkBox1klassClick)
    tabWidgetPeriods.currentChanged.connect(lambda: changedDateEdit())
    listWidgetHolidays2.itemClicked.connect(lambda: listWidgetHolidays2ItemClicked())

def listWidgetHolidays2ItemClicked():
    for i in range(len(range_date_holidays)):
        range_date_holidays[i][2] = str(int(listWidgetHolidays2.item(i).checkState() == QtCore.Qt.CheckState.Checked))
    saveHolidays2ToFile()
    makeListOfDays()
    loadWeekDays()
    readWeekDays()

def checkBox1klassClick(state):
    global klass1
    if state:
        klass1 = True
        dateEdit8.show()
        label1klass.show()
        dateEdit9.show()
    else:
        klass1 = False
        dateEdit8.hide()
        label1klass.hide()
        dateEdit9.hide()
    makeListOfDays()
    loadWeekDays()
    readWeekDays()

def pushButtonDeleteReplaceClick():
    global dict_date_replaced
    listItems = win_dialog.listWidgetReplace.selectedItems()
    if not listItems:
        return
    for item in listItems:
        del dict_date_replaced[item.text().split()[0]]
    saveReplacedToFile()
    loadReplaceToApp()
    makeListOfDays()
    loadWeekDays()
    readWeekDays()

def pushButtonAddReplaceClick():
    global dict_date_replaced
    dict_date_replaced[win_dialog.dateEditReplace.dateTime().toString('dd.MM.yyyy')] = win_dialog.comboBoxReplace.currentIndex()
    saveReplacedToFile()
    loadReplaceToApp()
    makeListOfDays()
    loadWeekDays()
    readWeekDays()

def sort_comp(el):
    el = el[0].split('.')[::-1]
    return int(''.join(el))

def pushButtonAddHolidays2Click():
    global range_date_holidays
    range_date_holidays.append([dateEditAdd1.dateTime().toString('dd.MM.yyyy'), dateEditAdd2.dateTime().toString('dd.MM.yyyy'), '2'])
    range_date_holidays.sort(key=sort_comp)
    saveHolidays2ToFile()
    loadHolidays2ToApp()
    makeListOfDays()
    loadWeekDays()
    readWeekDays()

def pushButtonDeleteHolidayClick():
    global list_date_holidays
    listItems = listWidgetHolidays.selectedItems()
    if not listItems:
        return
    for item in listItems:
        list_date_holidays.remove(item.text())
    loadHolidaysToApp()
    saveHolidaysToFile()
    makeListOfDays()
    loadWeekDays()
    readWeekDays()

def pushButtonDeleteHoliday2Click():
    global range_date_holidays
    listItems = listWidgetHolidays2.selectedItems()
    if not listItems:
        return
    for item in listItems:
        range_date_holidays.remove(item.text().split()[::2] + [str(item.checkState())])
    loadHolidays2ToApp()
    saveHolidays2ToFile()
    makeListOfDays()
    loadWeekDays()
    readWeekDays()

def pushButtonAddHolidayClick():
    global list_date_holidays
    list_date_holidays.append(dateEditHoliday.dateTime().toString('dd.MM.yyyy'))
    loadHolidaysToApp()
    saveHolidaysToFile()
    makeListOfDays()
    loadWeekDays()
    readWeekDays()

def checkBoxDiffHoursClick(state):
    global diff_hours
    diff_hours = state
    if diff_hours:
        win.tabWidget.tabBar().show()
    else:
        win.tabWidget.setCurrentIndex(0)
        win.tabWidget.tabBar().hide()
    readWeekDays()

def checkBoxRewriteClick(state):
    global rewrite
    rewrite = state
    pushButtonFill.setText(pushButtonFill_text_dict[rewrite])


class CheckVersionThread(QtCore.QThread):
    labelStatusChangeSignal = QtCore.pyqtSignal(str)
    labelLinkShowSignal = QtCore.pyqtSignal(bool)
    def  __init__(self, parent=None):
        QtCore.QThread.__init__(self, parent)
    def run(self):
        self.labelStatusChangeSignal.emit('Проверка...')
        self.labelLinkShowSignal.emit(False)
        new_ver = checkNewVersion()
        if new_ver:
            if new_ver == -1:
                self.labelStatusChangeSignal.emit('Произошла ошибка')
            else:
                self.labelStatusChangeSignal.emit('Доступна новая версия '+new_ver+'!')
                self.labelLinkShowSignal.emit(True)
        else:
            self.labelStatusChangeSignal.emit('Вы используете актуальную версию')

def on_label_status_change(s):
    win_help.label_status.setText(s)
    win.labelNewVersion.setText(s)

def on_label_link_show(show):
    if show:
        win_help.groupBoxLink.show()
        win.groupBoxNewVer.show()
    else:
        win_help.groupBoxLink.hide()

check_version_thread = CheckVersionThread()
check_version_thread.labelStatusChangeSignal.connect(on_label_status_change)
check_version_thread.labelLinkShowSignal.connect(on_label_link_show)

def checkNewVersion():
    try:
        url = 'https://sourceforge.net/projects/autofillktp/files/'
        html = urlopen(url, timeout=5).read().decode('utf-8')
        s = re.findall(r'latest.*title=(.*?)\.zip', html)
        new_ver = s[0].split()[-1]
        old_ver = version.split()[1]
        if int(float(new_ver)*100) > int(float(old_ver)*100):
            return(new_ver)
    except:
        return(-1)
    return(False)

def buttonCheckNewVersionClick():
    check_version_thread.start()

def setTable(text):
    global table_number
    table_number = int(text.split()[-1])-1
    read_hours_thread.start()
    

def setColumnWithDay(index):
    if index <= 0:
        return
    global column_with_days
    column_with_days = index
    pushButtonFill.setText(pushButtonFill_text_dict[rewrite])
    pushButtonFill.setStyleSheet('QWidget {font-size: 16pt;}')
    for i in range(6):
        globals()['checkBoxWeek'+str(i)].setChecked(False)
        globals()['spinBoxWeek'+str(i)].setValue(1)
    readWeekDays()


def changedDateEdit():
    global list_days, dict_days, all_days
    for i in range(0, 15):
        if i in (7, 8, 9):
            continue
        d1 = globals()['dateEdit'+str(i)].dateTime().toString('dd.MM.yyyy')
        d2 = globals()['dateEdit'+str(i+1)].dateTime().toString('dd.MM.yyyy')

        if not checkValidDatesArr(d1, d2):
            globals()['dateEdit'+str(i)].setStyleSheet('QWidget {background-color: rgb(255, 190, 191);%s}'%WIN_css)
            globals()['dateEdit'+str(i+1)].setStyleSheet('QWidget {background-color: rgb(255, 190, 191);%s}'%WIN_css)
            list_days = []
            dict_days = {0: 0, 1: 0, 2: 0, 3: 0, 4: 0, 5: 0, 6: 0}
            all_days = 0
            loadWeekDays()
            readWeekDays(err=True)
            frame_2.setEnabled(False)
            return
        else:
            globals()['dateEdit'+str(i)].setStyleSheet('QWidget {%s}'%WIN_css)
            globals()['dateEdit'+str(i+1)].setStyleSheet('QWidget {%s}'%WIN_css)
            frame_2.setEnabled(True)
    saveDateToFile()
    makeListOfDays()
    loadWeekDays()
    readWeekDays()

def loadWeekDays():
    for i in range(6):
        globals()['weekday'+str(i)].setText(str(dict_days[i]))
    label_all_days.setText(str(all_days))

def saveDateToFile():
    global list_date
    list_date = [0] * 16
    for i in range(16):
        list_date[i] = globals()['dateEdit'+str(i)].dateTime().toString('dd.MM.yyyy')
    with open(os.path.join(APP_PATH, 'сохранённый диапазон дат.txt'), 'w') as f:
        f.write('\n'.join(list_date))
    print('Dates saved to file')

def saveHolidaysToFile():
    with open(os.path.join(APP_PATH, 'сохранённые праздничные дни.txt'), 'w') as f:
        f.write('\n'.join(list_date_holidays))
    print('Holidays saved to file')

def saveHolidays2ToFile():
    with open(os.path.join(APP_PATH, 'сохранённые дополнительные каникулы.txt'), 'w') as f:
        for el in range_date_holidays:
            f.write(el[0] + ' ' + el[1] + ' ' + str(el[2]) + '\n')
    print('Holidays2 saved to file')

def saveReplacedToFile():
    l = []
    for key, value in dict_date_replaced.items():
        l.append(key+';'+str(value))
    with open(os.path.join(APP_PATH, 'сохранённые замены дней недели.txt'), 'w') as f:
        f.write('\n'.join(l))
    print('Replaced saved to file')

def makeListOfDays(weekday=[], selected = False, doubleday=[], weekday2=[], doubleday2=[]):
    global list_days, all_days, dict_days, list_klass1
    if selected:
        list_days = []
    all_days = 0
    dict_days = {0: 0, 1: 0, 2: 0, 3: 0, 4: 0, 5: 0, 6: 0}
    list_all_days=(0, 1, 2, 3, 4, 5)
    year = ''
    if checkBoxYear.isChecked():
        year = '.%y'
    new_year = list_date[4].split('.')[-1]
    list_holidays = []
    # список дополнительных каникул
    for el in range_date_holidays:
        if el[2] == '0':
            continue
        d1 = datetime.strptime(el[0], '%d.%m.%Y') 
        d2 = datetime.strptime(el[1], '%d.%m.%Y')
        delta = d2 - d1
        if delta.days > 0:
            for i in range(delta.days + 1):
                d = d1 + timedelta(i)
                list_holidays.append(d.strftime('%d.%m.%Y'))
    list_klass1 = []
    # каникулы 1 класса
    d1 = datetime.strptime(list_date[8], '%d.%m.%Y') 
    d2 = datetime.strptime(list_date[9], '%d.%m.%Y')
    delta = d2-d1
    if delta.days > 0:
        for i in range(delta.days + 1):
            d = d1 + timedelta(i)
            list_klass1.append(d.strftime('%d.%m.%Y'))
    if tabWidgetPeriods.currentIndex() == 0:
        r = range(0, 7, 2)
    else:
        r = range(10, 15, 2)
    for i in r:
        d1 = datetime.strptime(list_date[i], '%d.%m.%Y')
        d2 = datetime.strptime(list_date[i+1], '%d.%m.%Y')
        delta = d2 - d1
        for j in range(delta.days + 1):
            d = d1 + timedelta(j)
            day = d.weekday()
            if d.strftime('%d.%m.%Y') in dict_date_replaced.keys():
                day = dict_date_replaced[d.strftime('%d.%m.%Y')]
            if d.strftime('%d.%m.%Y') not in list_date_holidays:
                if klass1:
                    if d.strftime('%d.%m.%Y') in list_klass1:
                        continue
                if d.strftime('%d.%m.%Y') in list_holidays:
                    continue
                if diff_hours and d.year == int(new_year):
                    weekday = weekday2
                    doubleday = doubleday2
                if day in weekday:
                    list_days.append(d.strftime('%d.%m'+year))
                    if day in doubleday:
                        list_days.append(d.strftime('%d.%m'+year))
                if day in list_all_days:
                    all_days += 1
                dict_days[day] += 1

def checkValidDatesArr(d1, d2):
    d1 = datetime.strptime(d1, '%d.%m.%Y')
    d2 = datetime.strptime(d2, '%d.%m.%Y')
    delta = d2 - d1
    if delta.days <=0:
        return False
    else:
        return True

def readWeekDays(err = False):
    global week_days, double_days, week_days2, double_days2
    week_days = []
    double_days = []
    week_days2 = []
    double_days2 = []
    frame_2.setStyleSheet('QWidget {}')
    pushButtonFill.setText(pushButtonFill_text_dict[rewrite])
    pushButtonFill.setStyleSheet('QWidget {font-size: 16pt;}')
    labelProgress.setText('Заполнено 0 из '+str(count_hours))
    progressBar.setValue(0)
    checkBoxSetCol.setChecked(False)
    for i in range(6):
        if globals()['checkBoxWeek'+str(i)].isChecked():
            week_days.append(i)
        if globals()['spinBoxWeek'+str(i)].value() == 2:
            double_days.append(i)
        if globals()['checkBoxWeek'+str(i)+'_2'].isChecked():
            week_days2.append(i)
        if globals()['spinBoxWeek'+str(i)+'_2'].value() == 2:
            double_days2.append(i)
    if not err:
        makeListOfDays(week_days, True, double_days, week_days2, double_days2)
    label_selected_days.setText('Выбрано: '+str(len(list_days)))
    plainTextEdit.setPlainText('\n'.join(list_days))
    if file_name:
        if count_hours == len(list_days) and count_hours != 0:
            label_selected_days.setStyleSheet('QWidget {color: rgb(28, 153, 0);%s}'%WIN_css)
        else:
            label_selected_days.setStyleSheet('QWidget {color: rgb(255, 0, 0);%s}'%WIN_css)

class ReadHoursThread(QtCore.QThread):
    labelHoursChangeSignal = QtCore.pyqtSignal(str)
    def  __init__(self, parent=None):
        QtCore.QThread.__init__(self, parent)
    def run(self):
        global count_hours, one, column_with_days, list_one
        table = getTable()
        count_hours = 0
        if table:
            one = 2
            find = False
            rows = table.rows
            for i in range(1, 7):
                row = rows[i].cells
                if find:
                    break
                for j in range(one, len(row)):
                    if row[j].text.strip() == '1':
                        one = j
                        column_with_days = one + 1
                        find = True
                        break
            list_one = []
            columns = table.columns[one].cells
            for cell in columns:
                list_one.append(0)
                if cell.text.strip() == '1':
                    list_one[-1] = 1
                    count_hours += 1
                    self.labelHoursChangeSignal.emit('Найдено часов: '+str(count_hours))
                    sleep(0.001)
            if 1 not in list_one or count_hours%17 != 0:
                count_hours = 0
                one = 0
                column_with_days = 1
                list_one = []
                columns = table.columns[one].cells
                for i in range(len(columns)):
                    list_one.append(0)
                    if columns[i].text.strip() != '':
                        if columns[i].text.strip()[0].isdigit():
                            # if columns[i].text.strip() != table.columns[one+1].cells[i].text.strip():
                            list_one[-1] = 1
                            count_hours += 1
                            # print(columns[i].text.strip())
                            self.labelHoursChangeSignal.emit('Найдено часов: '+str(count_hours))
                            sleep(0.001)
            self.labelHoursChangeSignal.emit('Найдено часов: '+str(count_hours))
        else:
            self.labelHoursChangeSignal.emit('Таблица не найдена')

def on_finished_read_hours():
    pushButtonFill.setEnabled(True)
    pushButton.setEnabled(True)
    comboBoxTables.setEnabled(True)
    comboBoxColumns.setEnabled(True)
    checkBoxSetCol.setChecked(False)
    labelProgress.setText('Заполнено 0 из '+str(count_hours))
    progressBar.setValue(0)
    if count_hours % 17 == 0 and count_hours != 0:
        labelHours.setStyleSheet('QWidget {color: rgb(28, 153, 0);%s}'%WIN_css)
    else:
        labelHours.setStyleSheet('QWidget {color: rgb(255, 0, 0);%s}'%WIN_css)
        labelHours.setText('Найдено часов: '+str(count_hours)+'\nПроверьте столбец "Количество часов"')
    if count_hours == len(list_days) and count_hours != 0:
        label_selected_days.setStyleSheet('QWidget {color: rgb(28, 153, 0);%s}'%WIN_css)
    else:
        label_selected_days.setStyleSheet('QWidget {color: rgb(255, 0, 0);%s}'%WIN_css)
    if table_number != -1:
        win.comboBoxColumns.addItems(getColumnsNames())
    win.comboBoxColumns.setCurrentIndex(column_with_days)
    if win.comboBoxTables.count() == 0:
        win.comboBoxTables.addItems(getListOfTables())
    pushButtonFill.setFocus()

def on_started_read_hours():
    pushButtonFill.setEnabled(False)
    pushButton.setEnabled(False)
    comboBoxTables.setEnabled(False)
    comboBoxColumns.setEnabled(False)
    labelHours.setStyleSheet('QWidget {%s}'%WIN_css)
    label_selected_days.setStyleSheet('QWidget {%s}'%WIN_css)
    labelHours.setText('Найдено часов: 0')
    win.comboBoxColumns.clear()
    win.resize(100, 100)

def on_label_hours_change(s):
    labelHours.setText(s)

read_hours_thread = ReadHoursThread()
read_hours_thread.finished.connect(on_finished_read_hours)
read_hours_thread.started.connect(on_started_read_hours)
read_hours_thread.labelHoursChangeSignal.connect(on_label_hours_change)

def openFiles():
    global file_name, table_number
    frame_3.setStyleSheet('QWidget {}')
    pushButtonFill.setText(pushButtonFill_text_dict[rewrite])
    pushButtonFill.setStyleSheet('QWidget {font-size: 16pt;}')
    label_files.setText('')
    progressBar.setValue(0)
    labelProgress.setText('')
    labelHours.setText('')
    label_selected_days.setStyleSheet('QWidget {%s}'%WIN_css)
    file_name = QtWidgets.QFileDialog.getOpenFileName(win, 'Выберите файл', '', '(*.docx)')[0]
    if file_name:
        label_files.setText(os.path.basename(file_name))
        table_number = -1
        win.comboBoxTables.clear()
        read_hours_thread.start()

class FillTableThread(QtCore.QThread):
    labelProgressChangeSignal = QtCore.pyqtSignal(str)
    progressBarChangeSignal = QtCore.pyqtSignal(int)
    def  __init__(self, parent=None):
        QtCore.QThread.__init__(self, parent)
    def run(self):
        doc = docx.Document(file_name)
        tables = doc.tables
        table = tables[table_number]
        if table:
            filled = 0
            col = table.columns[column_with_days].cells
            for i in range(len(col)):
                if list_one[i] == 1:
                    col[i].text = list_days[filled]
                    filled += 1
                    self.labelProgressChangeSignal.emit('Заполнено '+str(filled)+' из '+str(count_hours))
                    self.progressBarChangeSignal.emit(filled*100//count_hours)
                    sleep(0.001)
            if rewrite:
                doc.save(file_name)
            else:
                doc.save(file_name[:-5]+' заполнено.docx')

def on_label_progress_change(s):
    labelProgress.setText(s)

def on_rogress_bar_change(n):
    progressBar.setValue(n)

def on_finished_fill_table():
    for i in range(4):
        globals()['frame_'+str(i)].setEnabled(True)
    pushButtonFill.setEnabled(True)
    pushButtonFill.setText('Готово')
    pushButtonFill.setStyleSheet('QWidget {font-size: 18pt; color: rgb(28, 153, 0);}')

def on_started_fill_table():
    for i in range(4):
        globals()['frame_'+str(i)].setEnabled(False)
    pushButtonFill.setEnabled(False)

fill_table_thread = FillTableThread()
fill_table_thread.finished.connect(on_finished_fill_table)
fill_table_thread.started.connect(on_started_fill_table)
fill_table_thread.labelProgressChangeSignal.connect(on_label_progress_change)
fill_table_thread.progressBarChangeSignal.connect(on_rogress_bar_change)


def fill():
    global list_days
    pushButtonFill.setText(pushButtonFill_text_dict[rewrite])
    pushButtonFill.setStyleSheet('QWidget {font-size: 16pt;}')
    if len(list_days) == 0 or not file_name or abs(len(list_days)-count_hours) > 7 or count_hours == 0 or not checkBoxSetCol.isChecked():
        if len(list_days) == 0 or abs(len(list_days)-count_hours) > 7:
            frame_2.setStyleSheet('QWidget {background-color: rgb(255, 190, 191);}')
        if not file_name or count_hours == 0:
            frame_3.setStyleSheet('QWidget {background-color: rgb(255, 190, 191);}')
        if not checkBoxSetCol.isChecked():
            groupBoxSetCol.setStyleSheet('QWidget {background-color: rgb(255, 190, 191);}')
        return
    if count_hours == 0:
        return
    if len(list_days) < count_hours:
        list_days.append('')
    fill_table_thread.start()

def getTable():
    if table_number == -1:
        getListOfTables()
    if table_number == -1:
        return
    doc = docx.Document(file_name)
    tables = doc.tables
    return(tables[table_number])

def getListOfTables():
    global table_number
    list_tables = []
    doc = docx.Document(file_name)
    tables = doc.tables
    for i in range(len(tables)):
        if len(tables[i].rows) > 17:
            if table_number == -1:
                table_number = i
                return
            list_tables.append('Таблица '+str(i+1))
    return(list_tables)

def getColumnsNames():
    list_of_columns = []
    doc = docx.Document(file_name)
    tables = doc.tables
    line = 0
    rows = tables[table_number].rows
    line = rows[0].cells[0]._tc.bottom
    for i in range(len(rows[0].cells)):
        s = []
        for j in range(line):
            t = rows[j].cells[i].text.strip().replace('\n', ' ')
            if t not in s:
                s.append(t)
        list_of_columns.append(' > '.join(s))
    return(list_of_columns)

pushButtonFill.setFocus()
fixFontIfWIN()
loadDateFromFile()
loadHolidaysFromFile()
loadHolidays2FromFile()
loadReplacedFromFile()
loadDateToApp()
loadHolidaysToApp()
loadHolidays2ToApp()
loadReplaceToApp()
makeListOfDays()
loadWeekDays()
connectSignals()

checkBox1klassClick(False)
win.resize(100, 100)

win.show()
check_version_thread.start()
sys.exit(app.exec())
